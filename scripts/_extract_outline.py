"""Extract paragraphs from kospi200_mm_project_outline.docx for analysis."""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document

DOC = ROOT / "samples" / "kospi200_mm_project_outline.docx"
OUT = ROOT / "scripts" / "_outline_dump.txt"


def main():
    doc = Document(str(DOC))
    lines: list[str] = []
    lines.append(f"=== {DOC.name} ===")
    lines.append(f"paragraphs: {len(doc.paragraphs)}")
    lines.append(f"tables: {len(doc.tables)}")
    lines.append("")
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        style = p.style.name if p.style else "?"
        lines.append(f"[{i:03d}] ({style}) {p.text}")
    for ti, t in enumerate(doc.tables):
        lines.append(f"\n--- Table {ti} ({len(t.rows)} rows) ---")
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip().replace("\n", " | ")
                if txt:
                    lines.append(f"  T{ti} R{ri} C{ci}: {txt}")
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(f"wrote {len(lines)} lines -> {OUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
