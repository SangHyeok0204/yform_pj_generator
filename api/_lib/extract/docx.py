"""DOCX text extraction via python-docx."""
from __future__ import annotations

from io import BytesIO

from docx import Document


def extract_docx_text(data: bytes) -> str:
    """Return paragraph + table cell text from a .docx file."""
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)
